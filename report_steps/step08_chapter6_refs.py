# -*- coding: utf-8 -*-
"""
Chương 6: Kết luận và hướng phát triển + Tài liệu tham khảo.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_body


def build(doc):
    doc.add_heading("CHƯƠNG 6. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)

    # ── 6.1 ──────────────────────────────────
    doc.add_heading("6.1. Kết luận", level=2)

    add_body(
        doc,
        "Đề tài \"SmartFridge AI – Quản gia tủ lạnh thông minh\" đã được "
        "nhóm nghiên cứu, thiết kế và triển khai thành công một hệ thống "
        "ứng dụng di động toàn diện phục vụ nhu cầu quản lý thực phẩm "
        "thông minh cho các hộ gia đình. Hệ thống kết hợp nhiều công nghệ "
        "hiện đại từ React Native, Spring Boot, FastAPI cho đến Gemini AI "
        "và WebSocket, tạo nên một sản phẩm có tính thực tiễn cao."
    )
    add_body(
        doc,
        "Về phần cá nhân, Nguyễn Đình Phúc đã hoàn thành đầy đủ các module "
        "được giao bao gồm: hệ thống xác thực đa phương thức (email/mật "
        "khẩu và Google OAuth2), quản lý Household với cơ chế mời thành "
        "viên qua QR code, quản lý kho thực phẩm với logic FIFO đảm bảo "
        "hiệu quả sử dụng nguyên liệu, đồng bộ thời gian thực giữa nhiều "
        "thiết bị bằng WebSocket/STOMP, và giao diện người dùng cho các "
        "màn hình Login, Inventory, Quản lý thành viên và QR invitation."
    )
    add_body(
        doc,
        "Qua quá trình thực hiện đồ án, nhóm đã tích lũy được nhiều kinh "
        "nghiệm quý báu về việc làm việc nhóm theo mô hình Vertical Slice, "
        "tích hợp nhiều dịch vụ trong kiến trúc microservice, và ứng dụng "
        "trí tuệ nhân tạo vào bài toán thực tế. Sản phẩm cuối cùng đã đạt "
        "được mục tiêu là một trợ lý bếp thông minh, giúp người dùng tiết "
        "kiệm thời gian, giảm lãng phí thực phẩm và nâng cao trải nghiệm "
        "quản lý bếp núc hằng ngày."
    )

    # ── 6.2 ──────────────────────────────────
    doc.add_heading("6.2. Hạn chế", level=2)

    add_body(
        doc,
        "Mặc dù hệ thống đã hoạt động ổn định trong môi trường demo, vẫn "
        "còn một số hạn chế cần được nhận diện và cải thiện:"
    )
    add_body(
        doc,
        "Thứ nhất, độ chính xác của AI Vision và OCR chưa đạt mức tuyệt "
        "đối. Việc nhận diện thực phẩm từ hình ảnh tủ lạnh phụ thuộc nhiều "
        "vào điều kiện ánh sáng, góc chụp và độ phân giải ảnh. Hóa đơn "
        "viết tay hoặc hóa đơn bị nhạt cũng có thể gây sai sót trong quá "
        "trình trích xuất."
    )
    add_body(
        doc,
        "Thứ hai, Recommendation Engine cần một lượng dữ liệu công thức lớn "
        "hơn để phát huy tối ưu hiệu quả của tìm kiếm vector và RAG. Trong "
        "môi trường demo với số lượng công thức hạn chế, kết quả gợi ý có "
        "thể chưa đa dạng."
    )
    add_body(
        doc,
        "Thứ ba, chatbot ngữ cảnh hiện tại hoạt động dựa trên Gemini LLM, "
        "có thể sinh ra câu trả lời không chính xác 100% (hallucination). "
        "Cần thêm cơ chế kiểm tra fact-checking và giới hạn phạm vi trả lời."
    )
    add_body(
        doc,
        "Thứ tư, đồng bộ thời gian thực bằng WebSocket hoạt động tốt trong "
        "môi trường LAN nhưng chưa được tối ưu cho tình huống scale lớn "
        "(nhiều household, nhiều thiết bị đồng thời). Cần có cơ chế load "
        "balancing và message queue khi triển khai production."
    )
    add_body(
        doc,
        "Thứ năm, ứng dụng chưa hỗ trợ đầy đủ chế độ offline. Một số chức "
        "năng như xem kho thực phẩm có thể hoạt động offline nhờ "
        "redux-persist, nhưng các thao tác cần đồng bộ (thêm/sửa/xóa) đòi "
        "hỏi kết nối mạng."
    )

    # ── 6.3 ──────────────────────────────────
    doc.add_heading("6.3. Hướng phát triển", level=2)

    add_body(
        doc,
        "Trong tương lai, nhóm dự kiến phát triển hệ thống theo các hướng "
        "chính sau:"
    )
    add_body(
        doc,
        "Mở rộng khả năng AI: Tích hợp thêm mô hình nhận diện giọng nói "
        "tiếng Việt để hỗ trợ nhập liệu bằng giọng nói. Cải thiện độ chính "
        "xác của OCR bằng cách fine-tune mô hình cho hóa đơn Việt Nam. Xây "
        "dựng hệ thống học tập từ phản hồi người dùng (feedback loop) để "
        "cải thiện chất lượng gợi ý theo thời gian."
    )
    add_body(
        doc,
        "Tối ưu hiệu năng và khả năng mở rộng: Sử dụng message queue "
        "(RabbitMQ/Kafka) cho WebSocket broadcasting khi scale. Triển khai "
        "Kubernetes cho auto-scaling từng dịch vụ. Thêm CDN cho nội dung "
        "tĩnh và hình ảnh."
    )
    add_body(
        doc,
        "Tăng cường tính năng: Phát triển chế độ offline đầy đủ với sync "
        "queue. Tích hợp IoT cảm biến nhiệt độ tủ lạnh. Xây dựng community "
        "recipe sharing across households. Phát triển dashboard web cho quản "
        "lý trên máy tính."
    )
    add_body(
        doc,
        "Cải thiện trải nghiệm người dùng: Thêm animation và "
        "micro-interaction cho giao diện mượt hơn. Hỗ trợ đa ngôn ngữ "
        "(tiếng Anh, tiếng Việt). Tối ưu thời gian khởi động ứng dụng "
        "và tải dữ liệu."
    )

    doc.add_page_break()

    # ── Tài liệu tham khảo ──────────────────
    _references(doc)


def _references(doc):
    doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)

    refs = [
        "[1] Meta Platforms, \"React Native Documentation,\" 2024. [Online]. "
        "Available: https://reactnative.dev/docs/getting-started.",

        "[2] VMware, \"Spring Boot Reference Documentation,\" Version 3.1.x, "
        "2024. [Online]. Available: "
        "https://docs.spring.io/spring-boot/docs/current/reference/htmlsingle/.",

        "[3] S. Ramírez, \"FastAPI Documentation,\" 2024. [Online]. Available: "
        "https://fastapi.tiangolo.com/.",

        "[4] The PostgreSQL Global Development Group, \"PostgreSQL 15 "
        "Documentation,\" 2024. [Online]. Available: "
        "https://www.postgresql.org/docs/15/.",

        "[5] pgvector Contributors, \"pgvector: Open-source vector similarity "
        "search for Postgres,\" GitHub, 2024. [Online]. Available: "
        "https://github.com/pgvector/pgvector.",

        "[6] Google, \"Gemini API Documentation,\" 2024. [Online]. Available: "
        "https://ai.google.dev/docs.",

        "[7] Google Cloud, \"Cloud Vision API Documentation,\" 2024. [Online]. "
        "Available: https://cloud.google.com/vision/docs.",

        "[8] Redis Ltd., \"Redis Documentation,\" 2024. [Online]. Available: "
        "https://redis.io/docs/.",

        "[9] Firebase, \"Firebase Cloud Messaging Documentation,\" 2024. "
        "[Online]. Available: "
        "https://firebase.google.com/docs/cloud-messaging.",

        "[10] D. Hardt, \"The OAuth 2.0 Authorization Framework,\" RFC 6749, "
        "IETF, 2012.",

        "[11] M. Jones, J. Bradley, and N. Sakimura, \"JSON Web Token (JWT),\" "
        "RFC 7519, IETF, 2015.",

        "[12] I. Fette and A. Melnikov, \"The WebSocket Protocol,\" RFC 6455, "
        "IETF, 2011.",

        "[13] P. Lewis, E. Perez, A. Piktus et al., \"Retrieval-Augmented "
        "Generation for Knowledge-Intensive NLP Tasks,\" Advances in Neural "
        "Information Processing Systems, vol. 33, pp. 9459–9474, 2020.",

        "[14] R. Smith, \"An Overview of the Tesseract OCR Engine,\" Proc. "
        "Ninth Int. Conference on Document Analysis and Recognition (ICDAR), "
        "IEEE, 2007.",

        "[15] Redux Team, \"Redux Toolkit Documentation,\" 2024. [Online]. "
        "Available: https://redux-toolkit.js.org/.",

        "[16] Docker Inc., \"Docker Compose Documentation,\" 2024. [Online]. "
        "Available: https://docs.docker.com/compose/.",

        "[17] Flyway Contributors, \"Flyway Database Migrations,\" 2024. "
        "[Online]. Available: https://flywaydb.org/documentation/.",

        "[18] R. Johnson, J. Hoeller et al., \"Spring Framework Reference "
        "Documentation,\" VMware, 2024.",
    ]

    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        p.paragraph_format.space_after = Pt(4)

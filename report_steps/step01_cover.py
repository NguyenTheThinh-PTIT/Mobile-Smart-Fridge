# -*- coding: utf-8 -*-
"""
Step 1: Trang bìa báo cáo.
"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_centered, add_body


def build(doc):
    # Khoảng trống phía trên
    for _ in range(3):
        doc.add_paragraph()

    add_centered(doc, "TRƯỜNG ĐẠI HỌC CÔNG NGHỆ THÔNG TIN", size=16, bold=True)
    add_centered(doc, "KHOA CÔNG NGHỆ PHẦN MỀM", size=14, bold=True)
    add_centered(doc, "──────────", size=14)

    for _ in range(2):
        doc.add_paragraph()

    add_centered(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", size=20, bold=True, space_after=12)
    add_centered(doc, "ĐỀ TÀI:", size=16, bold=True)
    add_centered(doc, "SMARTFRIDGE AI", size=22, bold=True)
    add_centered(doc, "QUẢN GIA TỦ LẠNH THÔNG MINH", size=18, bold=True)

    for _ in range(2):
        doc.add_paragraph()

    add_body(doc, "Nhóm QLĐT: ................", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(doc, "Nhóm BTL: ................", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT)

    doc.add_paragraph()
    add_body(doc, "Thành viên nhóm:", bold=True,
             alignment=WD_ALIGN_PARAGRAPH.LEFT)

    members = [
        ("Đinh Viết Quang",  "................"),
        ("Nguyễn Đình Phúc", "................"),
        ("Tạ Trường Vũ",     "................"),
        ("Nguyễn Thế Thịnh", "................"),
    ]
    for name, mssv in members:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(2)
        run = p.add_run(f"{name}  —  MSSV: {mssv}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    doc.add_paragraph()
    add_body(doc,
             "Thành viên thực hiện báo cáo (cá nhân): Nguyễn Đình Phúc",
             bold=True, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    for _ in range(3):
        doc.add_paragraph()
    add_centered(doc, "TP. Hồ Chí Minh, 2025", size=14, bold=True)

    doc.add_page_break()

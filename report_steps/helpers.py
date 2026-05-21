# -*- coding: utf-8 -*-
"""
Helpers và cấu hình styles dùng chung cho toàn bộ báo cáo.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def create_document():
    """Tạo Document và cấu hình styles toàn cục."""
    doc = Document()

    # --- Normal style ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.space_before = Pt(0)

    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(
        f'<w:rFonts {nsdecls("w")} w:eastAsia="Times New Roman"/>'
    )
    rPr.append(rFonts)

    # --- Heading styles ---
    cfg = {
        1: (Pt(18), Pt(24), Pt(12)),
        2: (Pt(16), Pt(18), Pt(6)),
        3: (Pt(15), Pt(12), Pt(6)),
        4: (Pt(14), Pt(6),  Pt(6)),
    }
    for level, (sz, sb, sa) in cfg.items():
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.color.rgb = RGBColor(0, 0, 0)
        hs.font.bold = True
        hs.font.size = sz
        hs.paragraph_format.space_before = sb
        hs.paragraph_format.space_after = sa

    # --- Page setup ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.5)
        section.right_margin = Cm(2)
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)

    return doc


# ============================================================
# Hàm tiện ích dùng chung
# ============================================================

def add_centered(doc, text, size=14, bold=False, space_after=6, space_before=0):
    """Thêm đoạn văn bản căn giữa."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    return p


def add_body(doc, text, bold=False, italic=False,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Thêm đoạn văn bản nội dung."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = bold
    run.italic = italic
    return p


def _cell_text(cell, text, bold=False, size=12,
               alignment=WD_ALIGN_PARAGRAPH.LEFT):
    """Ghi text vào ô bảng."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold


def _shade_row(row, color="D9E2F3"):
    """Tô nền cho hàng header bảng."""
    for cell in row.cells:
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading)


def make_table(doc, headers, rows, col_widths=None):
    """Tạo bảng với header và dữ liệu."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        _cell_text(table.rows[0].cells[i], h, bold=True, size=12,
                   alignment=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_row(table.rows[0])

    # Data
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            _cell_text(table.rows[r_idx + 1].cells[c_idx], str(val), size=12)

    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # khoảng trống sau bảng
    return table

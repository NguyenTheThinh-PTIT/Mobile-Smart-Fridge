# -*- coding: utf-8 -*-
"""
Chương 2: Phân tích yêu cầu hệ thống.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_centered, add_body, make_table, _cell_text


def build(doc):
    doc.add_heading("CHƯƠNG 2. PHÂN TÍCH YÊU CẦU HỆ THỐNG", level=1)

    # ── 2.1 ──────────────────────────────────
    doc.add_heading("2.1. Mô tả bài toán", level=2)

    add_body(
        doc,
        "Bài toán đặt ra là xây dựng một ứng dụng di động cho phép người dùng "
        "và gia đình họ quản lý toàn bộ thực phẩm trong tủ lạnh một cách thông "
        "minh và tự động. Hệ thống cần giải quyết các vấn đề cụ thể sau: "
        "(1) Lưu trữ và theo dõi trạng thái của từng loại thực phẩm trong kho, "
        "bao gồm số lượng, vị trí, ngày nhập và hạn sử dụng; "
        "(2) Đồng bộ dữ liệu giữa nhiều thiết bị của các thành viên trong cùng "
        "một gia đình theo thời gian thực; "
        "(3) Tự động nhận diện thực phẩm từ hình ảnh và hóa đơn mua sắm; "
        "(4) Gợi ý công thức nấu ăn phù hợp dựa trên nguyên liệu hiện có và "
        "sở thích cá nhân; "
        "(5) Lập kế hoạch bữa ăn và theo dõi dinh dưỡng cho cả gia đình."
    )
    add_body(
        doc,
        "Hệ thống hướng tới đối tượng sử dụng là các hộ gia đình Việt Nam, "
        "đặc biệt là những gia đình có nhiều thành viên cùng sử dụng chung "
        "bếp. Ứng dụng cần dễ sử dụng, hỗ trợ tiếng Việt và phù hợp với "
        "thói quen nấu ăn của người Việt."
    )

    # ── 2.2 ──────────────────────────────────
    doc.add_heading("2.2. Các tác nhân (Actors)", level=2)
    add_body(doc, "Hệ thống SmartFridge AI xác định các tác nhân sau:")

    actors = [
        ("Người dùng cá nhân",
         "Người sử dụng chính của ứng dụng, có thể đăng ký tài khoản, đăng "
         "nhập, quản lý thực phẩm cá nhân, xem gợi ý món ăn và sử dụng chatbot."),
        ("Thành viên gia đình (Member)",
         "Người được mời vào một Household, có quyền xem và cập nhật kho "
         "thực phẩm chung, xem kế hoạch bữa ăn."),
        ("Quản trị viên Household (Owner/Admin)",
         "Người tạo Household hoặc được cấp quyền Admin, có thêm quyền "
         "mời/xóa thành viên, chuyển quyền sở hữu, quản lý cài đặt nhà."),
        ("AI Services",
         "Bao gồm Gemini API (embedding và LLM), Google Vision API (OCR), "
         "hoạt động như tác nhân xử lý phía sau để cung cấp tính năng thông minh."),
        ("Notification Service (FCM)",
         "Dịch vụ thông báo đẩy của Firebase, gửi cảnh báo hết hạn, nhắc nhở "
         "kế hoạch nấu ăn và tổng kết tuần đến thiết bị người dùng."),
    ]
    for name, desc in actors:
        add_body(doc, f"{name}: {desc}")

    # ── 2.3 ──────────────────────────────────
    doc.add_heading("2.3. Yêu cầu chức năng", level=2)
    add_body(
        doc,
        "Các yêu cầu chức năng được nhóm theo bốn nhóm chính tương ứng với "
        "các vertical slice trong phân công nhóm:"
    )

    # Core & Sync
    add_body(doc, "Nhóm chức năng cơ bản (Core & Sync — Nguyễn Đình Phúc):", bold=True)
    make_table(doc, ["Mã", "Mô tả chức năng"], [
        ["FC01", "Đăng ký tài khoản bằng email/mật khẩu"],
        ["FC02", "Đăng nhập bằng email/mật khẩu hoặc Google OAuth2"],
        ["FC03", "Tạo Household (\"nhà\") mới"],
        ["FC04", "Mời thành viên vào Household bằng mã QR hoặc mã mời (6 ký tự)"],
        ["FC05", "Phân quyền Owner/Admin/Member trong Household"],
        ["FC06", "Xem, thêm, sửa, xóa thực phẩm trong kho (Inventory CRUD)"],
        ["FC07", "Quản lý nhiều lô (batch) cho cùng một loại thực phẩm với FIFO"],
        ["FC08", "Đồng bộ kho thực phẩm thời gian thực giữa các thiết bị qua WebSocket"],
        ["FC09", "Cập nhật hồ sơ cá nhân (profile CRUD)"],
        ["FC10", "Xem danh sách thành viên Household và quản lý quyền"],
    ])

    # Input & Shopping
    add_body(doc, "Nhóm chức năng AI đầu vào (Input & Shopping — Đinh Viết Quang):", bold=True)
    make_table(doc, ["Mã", "Mô tả chức năng"], [
        ["FA01", "Chụp ảnh tủ lạnh và nhận diện thực phẩm tự động (AI Vision)"],
        ["FA02", "Quét hóa đơn mua sắm và trích xuất thông tin (OCR)"],
        ["FA03", "Tạo và quản lý danh sách mua sắm (Shopping List)"],
        ["FA04", "Hỗ trợ chế độ offline cho scan và nhập liệu"],
    ])

    # Intelligence & Planning
    add_body(doc, "Nhóm chức năng thông minh (Intelligence & Planning — Tạ Trường Vũ):", bold=True)
    make_table(doc, ["Mã", "Mô tả chức năng"], [
        ["FI01", "Gợi ý công thức món ăn dựa trên nguyên liệu, sở thích và thời tiết"],
        ["FI02", "Lập kế hoạch bữa ăn theo tuần với lịch và điểm danh thành viên"],
        ["FI03", "Tính toán TDEE/BMR và ngân sách calo"],
        ["FI04", "Quản lý công thức (thêm, sửa, xóa, chia sẻ)"],
    ])

    # Assistant & Engagement
    add_body(doc, "Nhóm chức năng tương tác (Assistant & Engagement — Nguyễn Thế Thịnh):", bold=True)
    make_table(doc, ["Mã", "Mô tả chức năng"], [
        ["FE01", "Chatbot hỗ trợ hỏi đáp về dinh dưỡng và gợi ý nấu ăn"],
        ["FE02", "Tổng kết tuần (weekly recap) về tình hình sử dụng thực phẩm"],
        ["FE03", "Thông báo đẩy khi thực phẩm sắp hết hạn"],
        ["FE04", "Chế độ nấu ăn (cooking mode) với hướng dẫn từng bước"],
    ])

    # ── 2.4 ──────────────────────────────────
    doc.add_heading("2.4. Yêu cầu phi chức năng", level=2)

    make_table(doc, ["Mã", "Loại", "Mô tả"], [
        ["NFR01", "Hiệu năng",
         "Ứng dụng phải phản hồi trong vòng 2 giây cho các thao tác thông "
         "thường. API gợi ý món ăn phản hồi trong vòng 5 giây."],
        ["NFR02", "Dễ sử dụng",
         "Giao diện trực quan, hỗ trợ tiếng Việt, thao tác tối thiểu để "
         "hoàn thành nhiệm vụ."],
        ["NFR03", "Ổn định",
         "Hệ thống hoạt động ổn định, tự động xử lý lỗi và retry khi gặp "
         "sự cố mạng."],
        ["NFR04", "Đồng bộ dữ liệu",
         "Dữ liệu kho thực phẩm phải được đồng bộ real-time giữa các thiết "
         "bị trong cùng Household trong vòng 1 giây."],
        ["NFR05", "Bảo mật",
         "Xác thực bằng JWT với token hết hạn (30 phút), refresh token "
         "(30 ngày). Mật khẩu được hash bằng bcrypt."],
        ["NFR06", "Khả năng mở rộng",
         "Kiến trúc microservice cho phép scale từng thành phần độc lập. "
         "AI service tách biệt khỏi backend chính."],
        ["NFR07", "Bảo trì",
         "Mã nguồn tổ chức theo feature-driven structure, sử dụng "
         "TypeScript strict mode và design patterns rõ ràng."],
    ])

    # ── 2.5 ──────────────────────────────────
    doc.add_heading("2.5. Use Case tổng quan", level=2)

    add_body(
        doc,
        "Sơ đồ Use Case tổng quan thể hiện các nhóm chức năng chính của hệ "
        "thống SmartFridge AI và mối quan hệ giữa các tác nhân với hệ thống. "
        "Hệ thống gồm năm nhóm chức năng lớn: (1) Xác thực và quản lý tài "
        "khoản, (2) Quản lý Household và thành viên, (3) Quản lý kho thực "
        "phẩm, (4) Gợi ý và kế hoạch bữa ăn, (5) Chatbot và thông báo."
    )

    add_centered(doc,
                 "[Hình 2.1. Sơ đồ Use Case tổng quan hệ thống SmartFridge AI]",
                 size=12, bold=True)

    add_body(
        doc,
        "Ghi chú: Sơ đồ Use Case tổng quan bao gồm các tác nhân chính "
        "(Người dùng, Admin Household, AI Service, FCM) và các nhóm chức "
        "năng tương ứng. Người dùng tương tác trực tiếp với ứng dụng di "
        "động. Admin có thêm quyền quản lý thành viên. AI Service và FCM "
        "là các tác nhân hệ thống phía sau.",
        italic=True,
    )

    # ── 2.6 ──────────────────────────────────
    doc.add_heading("2.6. Use Case chi tiết", level=2)
    add_body(
        doc,
        "Phần này trình bày chi tiết các Use Case quan trọng thuộc phạm vi "
        "phụ trách của Nguyễn Đình Phúc."
    )

    _use_case_table(doc, "Bảng 2.3. Use Case Đăng nhập", [
        ["Tên Use Case",   "Đăng nhập hệ thống"],
        ["Tác nhân",       "Người dùng"],
        ["Mô tả",         "Người dùng nhập email và mật khẩu hoặc chọn đăng nhập "
                           "bằng Google OAuth2 để truy cập hệ thống"],
        ["Điều kiện trước", "Người dùng đã có tài khoản trong hệ thống"],
        ["Luồng chính",
         "1. Người dùng mở ứng dụng\n"
         "2. Hệ thống hiển thị màn hình Login\n"
         "3. Người dùng nhập email/mật khẩu\n"
         "4. Hệ thống xác thực thông tin\n"
         "5. Hệ thống trả về JWT (access token + refresh token)\n"
         "6. Ứng dụng lưu token và chuyển sang màn hình Home"],
        ["Luồng thay thế",
         "3a. Người dùng chọn Google OAuth2 → Hệ thống chuyển hướng sang "
         "Google → Google trả về token → Hệ thống xác thực và tạo JWT"],
        ["Điều kiện sau",  "Người dùng được xác thực và có thể sử dụng ứng dụng"],
        ["Ngoại lệ",      "Sai mật khẩu: hiện thông báo lỗi. Mạng lỗi: hiện "
                           "thông báo và retry."],
    ])

    _use_case_table(doc, "Bảng 2.4. Use Case Tạo Household", [
        ["Tên Use Case",   "Tạo Household mới"],
        ["Tác nhân",       "Người dùng đã đăng nhập"],
        ["Mô tả",         "Người dùng tạo một \"nhà\" (Household) mới để quản lý "
                           "thực phẩm chung với gia đình"],
        ["Luồng chính",
         "1. Người dùng chọn \"Tạo nhà mới\"\n"
         "2. Hệ thống yêu cầu nhập tên nhà\n"
         "3. Người dùng nhập tên và xác nhận\n"
         "4. Hệ thống tạo Household, gán người dùng làm Owner\n"
         "5. Hệ thống tự động tạo Inventory cho Household\n"
         "6. Hệ thống tạo mã mời (6 ký tự) và QR code\n"
         "7. Chuyển sang màn hình Household Overview"],
        ["Điều kiện sau",
         "Household được tạo với người dùng là Owner, có mã mời và QR "
         "để chia sẻ"],
    ])

    _use_case_table(doc, "Bảng 2.5. Use Case Quản lý kho thực phẩm", [
        ["Tên Use Case",   "Quản lý kho thực phẩm (Inventory CRUD)"],
        ["Tác nhân",       "Thành viên Household"],
        ["Mô tả",         "Thành viên xem, thêm, sửa, xóa thực phẩm trong kho "
                           "chung của Household"],
        ["Luồng chính",
         "1. Người dùng mở màn hình Inventory\n"
         "2. Hệ thống hiển thị danh sách thực phẩm hiện có\n"
         "3. Người dùng chọn thêm thực phẩm mới\n"
         "4. Nhập thông tin: tên, số lượng, đơn vị, vị trí lưu trữ, hạn sử dụng\n"
         "5. Hệ thống tạo inventory_batch mới với entry_date = ngày hiện tại\n"
         "6. Hệ thống phát sự kiện đồng bộ qua WebSocket\n"
         "7. Tất cả thiết bị của thành viên nhận cập nhật tức thì"],
        ["Luồng thay thế",
         "6a. Khi xóa thực phẩm: hệ thống áp dụng FIFO, ưu tiên trừ lô "
         "cũ nhất trước"],
        ["Điều kiện sau",
         "Kho thực phẩm được cập nhật và đồng bộ trên tất cả thiết bị"],
    ])

    doc.add_page_break()


def _use_case_table(doc, title, rows):
    """Tạo bảng Use Case dạng key-value."""
    add_body(doc, title, bold=True)
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(rows):
        _cell_text(t.rows[i].cells[0], k, bold=True, size=12)
        _cell_text(t.rows[i].cells[1], v, size=12)
    doc.add_paragraph()

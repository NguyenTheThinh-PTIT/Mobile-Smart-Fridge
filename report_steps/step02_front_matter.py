# -*- coding: utf-8 -*-
"""
Step 2: Trang thành viên, mục lục, danh sách viết tắt, hình vẽ, bảng biểu.
"""
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_centered, add_body, make_table


def build(doc):
    _members_page(doc)
    _table_of_contents(doc)
    _abbreviations(doc)
    _figures_list(doc)
    _tables_list(doc)


# ─────────────────────────────────────────────
def _members_page(doc):
    doc.add_heading("THÀNH VIÊN THAM GIA VÀ ĐÓNG GÓP", level=1)

    add_body(
        doc,
        "Bảng dưới đây trình bày chi tiết phân công nhiệm vụ và đóng góp của "
        "từng thành viên trong nhóm phát triển ứng dụng SmartFridge AI. Mô hình "
        "phân công được thực hiện theo kiểu Vertical Slice, theo đó mỗi thành viên "
        "chịu trách nhiệm toàn bộ một \"lát dọc\" của hệ thống, từ giao diện người "
        "dùng, logic nghiệp vụ cho đến tích hợp backend và cơ sở dữ liệu tương ứng. "
        "Cách chia này đảm bảo tính độc lập giữa các module, giúp các thành viên làm "
        "việc song song hiệu quả và giảm thiểu xung đột khi tích hợp."
    )

    headers = ["STT", "Họ tên", "MSSV", "Vai trò / Module",
               "Công nghệ chính", "Giao diện phụ trách", "Đóng góp"]
    rows = [
        ["1", "Đinh Viết Quang", "............",
         "Đầu vào & Mua sắm\n(Input & Shopping)",
         "AI Vision/OCR, Gemini API,\nReact Native Camera,\nOffline Storage",
         "Màn hình scan/crop,\nxác nhận kết quả scan,\nshopping list",
         "25%"],
        ["2", "Nguyễn Đình Phúc", "............",
         "Hạt nhân & Đồng bộ\n(Core & Sync)",
         "WebSocket/STOMP,\nSpring Security, JWT,\nGoogle OAuth2,\nRedux Toolkit,\nReact Navigation, QR Code",
         "Splash/Login, Inventory,\nQuản lý thành viên,\nQR invitation",
         "25%"],
        ["3", "Tạ Trường Vũ", "............",
         "Bộ não & Dinh dưỡng\n(Intelligence & Planning)",
         "FastAPI, Gemini Embedding,\npgvector, RAG pipeline,\nTDEE/BMR",
         "Recipe detail,\ngợi ý món ăn,\nlịch tuần (Planner)",
         "25%"],
        ["4", "Nguyễn Thế Thịnh", "............",
         "Trợ lý & Tương tác\n(Assistant & Engagement)",
         "LangChain4j, Gemini LLM,\nFCM, Batch Job, Redis",
         "Cooking mode,\nrecap tuần,\nfeed/chatbot",
         "25%"],
    ]
    make_table(doc, headers, rows)

    add_body(
        doc,
        "Ghi chú: Mô hình Vertical Slice đảm bảo mỗi thành viên sở hữu toàn bộ "
        "chiều dọc của một nhóm tính năng, từ frontend đến backend, bao gồm cả "
        "thiết kế cơ sở dữ liệu và logic nghiệp vụ. Điều này giúp từng cá nhân "
        "hiểu sâu về phần mình phụ trách và chịu trách nhiệm hoàn toàn với chất "
        "lượng sản phẩm của module đó."
    )
    doc.add_page_break()


# ─────────────────────────────────────────────
def _table_of_contents(doc):
    doc.add_heading("MỤC LỤC", level=1)

    add_body(doc,
             "[Mục lục sẽ được cập nhật tự động trong Microsoft Word: "
             "References → Table of Contents → Update Table]",
             italic=True)
    doc.add_paragraph()

    items = [
        ("Thành viên tham gia và đóng góp", 0),
        ("Mục lục", 0),
        ("Danh sách từ viết tắt", 0),
        ("Danh sách hình vẽ", 0),
        ("Danh sách bảng biểu", 0),
        ("Chương 1. Tổng quan về đề tài", 0),
        ("1.1. Lý do chọn đề tài", 1),
        ("1.2. Mục tiêu của ứng dụng", 1),
        ("1.3. Phạm vi nghiên cứu", 1),
        ("1.4. Concept của ứng dụng", 1),
        ("1.5. Ý nghĩa thực tiễn của đề tài", 1),
        ("1.6. Cấu trúc báo cáo", 1),
        ("Chương 2. Phân tích yêu cầu hệ thống", 0),
        ("2.1. Mô tả bài toán", 1),
        ("2.2. Các tác nhân (Actors)", 1),
        ("2.3. Yêu cầu chức năng", 1),
        ("2.4. Yêu cầu phi chức năng", 1),
        ("2.5. Use Case tổng quan", 1),
        ("2.6. Use Case chi tiết", 1),
        ("Chương 3. Phân tích và thiết kế hệ thống", 0),
        ("3.1. Kiến trúc tổng quan", 1),
        ("3.2. Thiết kế tổng quan ứng dụng", 1),
        ("3.3. Phân tích thiết kế chi tiết mức cá nhân của Nguyễn Đình Phúc", 1),
        ("3.4. Biểu đồ lớp", 1),
        ("3.5. Biểu đồ tuần tự", 1),
        ("3.6. Sơ đồ ERD", 1),
        ("3.7. Thiết kế giao diện", 1),
        ("Chương 4. Công nghệ và công cụ phát triển", 0),
        ("4.1. Công nghệ frontend", 1),
        ("4.2. Công nghệ backend", 1),
        ("4.3. Công nghệ AI service", 1),
        ("4.4. Cơ sở dữ liệu và vector database", 1),
        ("4.5. Thư viện và API bên ngoài", 1),
        ("4.6. Công cụ phát triển", 1),
        ("Chương 5. Cài đặt, triển khai và kết quả thực hiện", 0),
        ("5.1. Mô hình triển khai ứng dụng", 1),
        ("5.2. Các bước cài đặt và triển khai", 1),
        ("5.3. Kết quả thực hiện được", 1),
        ("5.4. Kết quả thử nghiệm/triển khai", 1),
        ("5.5. Đánh giá kết quả", 1),
        ("Chương 6. Kết luận và hướng phát triển", 0),
        ("6.1. Kết luận", 1),
        ("6.2. Hạn chế", 1),
        ("6.3. Hướng phát triển", 1),
        ("Tài liệu tham khảo", 0),
    ]
    for text, indent in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(indent)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

    doc.add_page_break()


# ─────────────────────────────────────────────
def _abbreviations(doc):
    doc.add_heading("DANH SÁCH TỪ VIẾT TẮT", level=1)

    data = [
        ("AI",        "Artificial Intelligence — Trí tuệ nhân tạo"),
        ("API",       "Application Programming Interface — Giao diện lập trình ứng dụng"),
        ("BMR",       "Basal Metabolic Rate — Tỷ lệ trao đổi chất cơ bản"),
        ("CRUD",      "Create, Read, Update, Delete — Các thao tác cơ bản trên dữ liệu"),
        ("DB",        "Database — Cơ sở dữ liệu"),
        ("ERD",       "Entity-Relationship Diagram — Sơ đồ thực thể – quan hệ"),
        ("FCM",       "Firebase Cloud Messaging — Dịch vụ thông báo đẩy của Firebase"),
        ("FIFO",      "First In First Out — Nhập trước xuất trước"),
        ("JWT",       "JSON Web Token — Mã thông báo xác thực dạng JSON"),
        ("LLM",       "Large Language Model — Mô hình ngôn ngữ lớn"),
        ("OAuth2",    "Open Authorization 2.0 — Giao thức ủy quyền mở phiên bản 2"),
        ("OCR",       "Optical Character Recognition — Nhận dạng ký tự quang học"),
        ("RAG",       "Retrieval-Augmented Generation — Sinh nội dung tăng cường truy xuất"),
        ("REST",      "Representational State Transfer — Kiến trúc truyền trạng thái"),
        ("STOMP",     "Simple Text Oriented Messaging Protocol"),
        ("TDEE",      "Total Daily Energy Expenditure — Tổng năng lượng tiêu hao hằng ngày"),
        ("UI",        "User Interface — Giao diện người dùng"),
        ("UX",        "User Experience — Trải nghiệm người dùng"),
        ("WebSocket", "Giao thức giao tiếp hai chiều theo thời gian thực"),
    ]
    make_table(doc, ["Viết tắt", "Giải thích"], [[a, b] for a, b in data])
    doc.add_page_break()


# ─────────────────────────────────────────────
def _figures_list(doc):
    doc.add_heading("DANH SÁCH HÌNH VẼ", level=1)

    figures = [
        ("Hình 2.1", "Sơ đồ Use Case tổng quan hệ thống SmartFridge AI"),
        ("Hình 3.1", "Kiến trúc tổng quan hệ thống SmartFridge AI"),
        ("Hình 3.2", "Kiến trúc nhiều lớp (multi-layer architecture)"),
        ("Hình 3.3", "Luồng xác thực JWT và Google OAuth2"),
        ("Hình 3.4", "Sơ đồ hoạt động quản lý Household"),
        ("Hình 3.5", "Luồng xử lý FIFO trong quản lý kho thực phẩm"),
        ("Hình 3.6", "Kiến trúc đồng bộ thời gian thực bằng WebSocket/STOMP"),
        ("Hình 3.7", "Biểu đồ lớp các thực thể chính"),
        ("Hình 3.8", "Biểu đồ tuần tự — Đăng nhập"),
        ("Hình 3.9", "Biểu đồ tuần tự — Tạo Household"),
        ("Hình 3.10", "Biểu đồ tuần tự — Quét QR tham gia Household"),
        ("Hình 3.11", "Biểu đồ tuần tự — Cập nhật kho và đồng bộ realtime"),
        ("Hình 3.12", "Sơ đồ ERD cơ sở dữ liệu SmartFridge AI"),
        ("Hình 3.13", "Thiết kế giao diện Login và Splash"),
        ("Hình 3.14", "Thiết kế giao diện Inventory"),
        ("Hình 3.15", "Thiết kế giao diện Quản lý thành viên và QR"),
        ("Hình 5.1", "Mô hình triển khai Docker Compose"),
        ("Hình 5.2", "Kết quả giao diện màn hình Login"),
        ("Hình 5.3", "Kết quả giao diện màn hình Inventory"),
        ("Hình 5.4", "Kết quả giao diện Household và QR"),
    ]
    make_table(doc, ["Ký hiệu", "Tên hình"], [[a, b] for a, b in figures])
    doc.add_page_break()


# ─────────────────────────────────────────────
def _tables_list(doc):
    doc.add_heading("DANH SÁCH BẢNG BIỂU", level=1)

    tables = [
        ("Bảng 1.1", "Phân công công việc theo Vertical Slice"),
        ("Bảng 2.1", "Danh sách yêu cầu chức năng"),
        ("Bảng 2.2", "Danh sách yêu cầu phi chức năng"),
        ("Bảng 2.3", "Mô tả Use Case Đăng nhập"),
        ("Bảng 2.4", "Mô tả Use Case Tạo Household"),
        ("Bảng 2.5", "Mô tả Use Case Quản lý kho thực phẩm"),
        ("Bảng 3.1", "Danh sách các lớp chính trong hệ thống"),
        ("Bảng 3.2", "Danh sách bảng cơ sở dữ liệu"),
        ("Bảng 4.1", "Công nghệ và công cụ sử dụng"),
        ("Bảng 5.1", "Kết quả thử nghiệm hệ thống"),
        ("Bảng 5.2", "Thống kê số liệu triển khai"),
    ]
    make_table(doc, ["Ký hiệu", "Tên bảng"], [[a, b] for a, b in tables])
    doc.add_page_break()

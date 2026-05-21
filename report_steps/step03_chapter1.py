# -*- coding: utf-8 -*-
"""
Chương 1: Tổng quan về đề tài.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_centered, add_body


def build(doc):
    doc.add_heading("CHƯƠNG 1. TỔNG QUAN VỀ ĐỀ TÀI", level=1)

    # ── 1.1 ──────────────────────────────────
    doc.add_heading("1.1. Lý do chọn đề tài", level=2)

    add_body(
        doc,
        "Trong bối cảnh xã hội hiện đại, lãng phí thực phẩm đang trở thành "
        "một vấn đề cấp bách mang tính toàn cầu. Theo báo cáo của Tổ chức "
        "Lương thực và Nông nghiệp Liên Hợp Quốc (FAO), mỗi năm trên thế "
        "giới có khoảng 1,3 tỷ tấn thực phẩm bị lãng phí, tương đương với "
        "một phần ba tổng sản lượng lương thực toàn cầu. Tại Việt Nam, tình "
        "trạng này cũng không kém phần nghiêm trọng khi nhiều gia đình "
        "thường xuyên phải bỏ đi thực phẩm hết hạn sử dụng hoặc bị hư hỏng "
        "do không kiểm soát được lượng tồn kho trong tủ lạnh."
    )
    add_body(
        doc,
        "Nguyên nhân chính của vấn đề này nằm ở việc thiếu một công cụ quản "
        "lý thực phẩm hiệu quả. Hầu hết các hộ gia đình Việt Nam vẫn quản "
        "lý tủ lạnh theo cách truyền thống: ghi nhớ bằng đầu hoặc đơn giản "
        "là không quản lý. Điều này dẫn đến nhiều hệ quả tiêu cực như: mua "
        "sắm trùng lặp, thực phẩm hết hạn không được phát hiện kịp thời, "
        "khó khăn trong việc lên kế hoạch bữa ăn hằng ngày và hằng tuần, "
        "cũng như sự bất hợp lý trong chế độ dinh dưỡng của các thành viên "
        "trong gia đình."
    )
    add_body(
        doc,
        "Đồng thời, sự phát triển vững chắc của công nghệ trí tuệ nhân tạo "
        "(AI), đặc biệt là các mô hình ngôn ngữ lớn (LLM) như Gemini, cũng "
        "như các dịch vụ nhận diện hình ảnh (Google Vision) và xử lý ngôn "
        "ngữ tự nhiên, đã mở ra những cơ hội mới để giải quyết bài toán này. "
        "AI có thể giúp nhận diện thực phẩm từ hình ảnh, phân tích hóa đơn "
        "mua sắm, gợi ý công thức nấu ăn dựa trên nguyên liệu hiện có, và "
        "cung cấp lời khuyên dinh dưỡng cá nhân hóa."
    )
    add_body(
        doc,
        "Ngoài ra, nhu cầu đồng bộ thông tin giữa các thành viên trong gia "
        "đình ngày càng tăng cao. Khi nhiều người cùng sử dụng chung một tủ "
        "lạnh, việc một người mua thêm thực phẩm mà người khác không biết, "
        "hoặc một người đã sử dụng hết nguyên liệu mà không cập nhật, là "
        "tình huống rất phổ biến. Do đó, một ứng dụng hỗ trợ đồng bộ thời "
        "gian thực giữa nhiều thiết bị sẽ mang lại giá trị thực tiễn to lớn."
    )
    add_body(
        doc,
        "Từ những lý do trên, nhóm chúng tôi quyết định thực hiện đề tài "
        "\"SmartFridge AI – Quản gia tủ lạnh thông minh\" với mục tiêu xây "
        "dựng một ứng dụng di động toàn diện, kết hợp quản lý kho thực phẩm, "
        "đồng bộ gia đình, trí tuệ nhân tạo nhận diện đầu vào, gợi ý món ăn, "
        "hỗ trợ nấu ăn và phân tích dinh dưỡng trong một nền tảng duy nhất."
    )

    # ── 1.2 ──────────────────────────────────
    doc.add_heading("1.2. Mục tiêu của ứng dụng", level=2)

    add_body(
        doc,
        "Mục tiêu tổng quát của đề tài là xây dựng một ứng dụng di động "
        "thông minh giúp người dùng quản lý thực phẩm trong tủ lạnh một "
        "cách hiệu quả, tự động hóa các quy trình nhập liệu, theo dõi hạn "
        "sử dụng, gợi ý công thức nấu ăn phù hợp và hỗ trợ lên kế hoạch "
        "bữa ăn cho cả gia đình.",
        bold=True,
    )

    add_body(doc, "Các mục tiêu cụ thể bao gồm:")

    objectives = [
        "Xây dựng hệ thống quản lý kho thực phẩm với cơ chế FIFO (nhập trước "
        "xuất trước), cho phép theo dõi chính xác số lượng, vị trí lưu trữ và "
        "hạn sử dụng của từng lô thực phẩm.",

        "Thiết kế và triển khai cơ chế đồng bộ thời gian thực bằng "
        "WebSocket/STOMP, đảm bảo mọi thay đổi về kho thực phẩm được cập nhật "
        "tức thì trên tất cả các thiết bị của thành viên trong gia đình.",

        "Xây dựng hệ thống xác thực và phân quyền đa tầng (Owner/Admin/Member) "
        "cho phép quản lý thành viên gia đình, mời thành viên qua mã QR và kiểm "
        "soát truy cập hợp lý.",

        "Tích hợp AI nhận diện thực phẩm từ hình ảnh tủ lạnh và hóa đơn mua "
        "sắm (OCR/Vision) giúp tự động hóa quá trình nhập liệu.",

        "Xây dựng engine gợi ý công thức món ăn thông minh dựa trên nguyên liệu "
        "hiện có, sở thích cá nhân, tình trạng dinh dưỡng và điều kiện thời tiết, "
        "sử dụng kỹ thuật RAG (Retrieval-Augmented Generation) với vector database.",

        "Phát triển chatbot ngôn ngữ tự nhiên hỗ trợ người dùng lên thực đơn, hỏi "
        "đáp về dinh dưỡng và nhận lời khuyên nấu ăn theo ngữ cảnh.",

        "Xây dựng chức năng lập kế hoạch bữa ăn (meal planner) với lịch tuần, "
        "theo dõi sự tham gia của từng thành viên và tính toán ngân sách calo.",

        "Triển khai hệ thống thông báo thông minh cảnh báo thực phẩm sắp hết hạn, "
        "nhắc nhở kế hoạch nấu ăn và tổng kết tuần.",
    ]
    for i, obj in enumerate(objectives, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"({i}) {obj}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    # ── 1.3 ──────────────────────────────────
    doc.add_heading("1.3. Phạm vi nghiên cứu", level=2)

    add_body(
        doc,
        "Phạm vi nền tảng: Ứng dụng được phát triển trên nền tảng di động "
        "(React Native) hỗ trợ cả Android và iOS, kết nối với backend Spring "
        "Boot và AI service FastAPI. Cơ sở dữ liệu sử dụng PostgreSQL với "
        "phần mở rộng pgvector cho tìm kiếm vector."
    )
    add_body(
        doc,
        "Phạm vi tính năng: Hệ thống tập trung vào các nhóm chức năng chính "
        "gồm quản lý kho thực phẩm, đồng bộ thời gian thực, xác thực và "
        "quản lý gia đình, gợi ý công thức, lập kế hoạch bữa ăn, chatbot "
        "và thông báo. Các tính năng nâng cao như phân tích video, tích hợp "
        "IoT cảm biến tủ lạnh nằm ngoài phạm vi của đề tài này."
    )
    add_body(
        doc,
        "Phạm vi AI: Ứng dụng sử dụng Gemini API cho cả embedding (3072 chiều) "
        "và sinh văn bản (LLM), Google Vision API cho nhận diện hình ảnh. Các "
        "mô hình AI được sử dụng thông qua API, không training từ đầu."
    )
    add_body(
        doc,
        "Phạm vi triển khai: Ứng dụng được triển khai trong môi trường Docker "
        "Compose gồm 4 dịch vụ (PostgreSQL, Spring Boot backend, FastAPI AI "
        "service, Redis), kiểm thử trên máy thật Android và emulator."
    )
    add_body(
        doc,
        "Giới hạn của đề tài: Ứng dụng hiện tại chưa hỗ trợ nhận diện giọng nói "
        "tiếng Việt chính xác 100%, chưa tích hợp IoT, và chất lượng gợi ý phụ "
        "thuộc vào độ phong phú của cơ sở dữ liệu công thức."
    )

    # ── 1.4 ──────────────────────────────────
    doc.add_heading("1.4. Concept của ứng dụng", level=2)

    add_body(
        doc,
        "SmartFridge AI được xây dựng theo concept của một trợ lý bếp thông minh "
        "toàn diện. Ứng dụng không đơn thuần là một công cụ quản lý danh sách "
        "thực phẩm, mà là một hệ sinh thái kết hợp nhiều thành phần công nghệ "
        "để phục vụ trọn vẹn nhu cầu của người dùng trong việc quản lý bếp núc "
        "hằng ngày."
    )
    add_body(
        doc,
        "Cụ thể, concept của ứng dụng bao gồm năm trụ cột chính: (1) Quản lý "
        "kho thực phẩm thông minh với FIFO và theo dõi hạn sử dụng; (2) Đồng "
        "bộ gia đình thời gian thực cho phép nhiều thành viên cùng quản lý chung "
        "một tủ lạnh; (3) AI nhận diện đầu vào giúp tự động hóa việc nhập liệu "
        "từ hình ảnh tủ lạnh hoặc hóa đơn; (4) Gợi ý món ăn và hỗ trợ nấu ăn "
        "dựa trên nguyên liệu thực tế, sở thích và tình trạng sức khỏe; (5) "
        "Phân tích dinh dưỡng và lập kế hoạch bữa ăn cho cả gia đình."
    )
    add_body(
        doc,
        "Điểm khác biệt của SmartFridge AI so với các ứng dụng quản lý thực phẩm "
        "thông thường là sự kết hợp chặt chẽ giữa quản lý dữ liệu thực phẩm "
        "(inventory management) với trí tuệ nhân tạo (AI-powered suggestions) và "
        "tính năng cộng tác gia đình (household collaboration). Người dùng không "
        "chỉ biết mình có gì trong tủ lạnh, mà còn được gợi ý nên nấu gì, khi "
        "nào nên nấu, và được cảnh báo khi thực phẩm sắp hết hạn — tất cả được "
        "đồng bộ giữa các thành viên gia đình theo thời gian thực."
    )

    # ── 1.5 ──────────────────────────────────
    doc.add_heading("1.5. Ý nghĩa thực tiễn của đề tài", level=2)

    add_body(
        doc,
        "Về mặt cá nhân và gia đình, ứng dụng giúp người dùng tiết kiệm thời "
        "gian và công sức trong việc quản lý thực phẩm hằng ngày. Thay vì phải "
        "nhớ hoặc ghi chép thủ công, người dùng chỉ cần chụp ảnh hoặc quét hóa "
        "đơn để cập nhật kho. Hệ thống cảnh báo hạn sử dụng giúp giảm thiểu "
        "lượng thực phẩm bị bỏ đi, qua đó tiết kiệm tài chính cho gia đình."
    )
    add_body(
        doc,
        "Về mặt sức khỏe, chức năng gợi ý món ăn dựa trên chỉ số dinh dưỡng "
        "(TDEE/BMR) và sở thích ăn uống giúp người dùng xây dựng chế độ ăn uống "
        "cân bằng hơn. Meal planner cho phép lên kế hoạch bữa ăn cả tuần, đảm "
        "bảo đa dạng và đủ chất dinh dưỡng."
    )
    add_body(
        doc,
        "Về mặt trải nghiệm người dùng, việc tích hợp chatbot và gợi ý thông "
        "minh tạo nên một trải nghiệm tương tác tự nhiên, giúp người dùng cảm "
        "thấy như có một trợ lý bếp riêng. Tính năng đồng bộ thời gian thực "
        "giải quyết triệt để vấn đề \"người này mua mà người kia không biết\" "
        "trong các gia đình nhiều thành viên."
    )

    # ── 1.6 ──────────────────────────────────
    doc.add_heading("1.6. Cấu trúc báo cáo", level=2)

    add_body(
        doc,
        "Báo cáo này gồm 6 chương chính: Chương 1 trình bày tổng quan về đề "
        "tài, lý do chọn đề tài và mục tiêu. Chương 2 phân tích yêu cầu hệ "
        "thống bao gồm yêu cầu chức năng và phi chức năng. Chương 3 là chương "
        "trọng tâm, trình bày phân tích và thiết kế hệ thống, bao gồm kiến "
        "trúc tổng quan và phân tích chi tiết phần cá nhân của Nguyễn Đình Phúc. "
        "Chương 4 mô tả công nghệ và công cụ phát triển được sử dụng. Chương 5 "
        "trình bày kết quả cài đặt, triển khai và thử nghiệm. Chương 6 kết luận "
        "và nêu hướng phát triển trong tương lai."
    )

    doc.add_page_break()

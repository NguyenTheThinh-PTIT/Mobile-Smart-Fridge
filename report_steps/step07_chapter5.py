# -*- coding: utf-8 -*-
"""
Chương 5: Cài đặt, triển khai và kết quả thực hiện.
"""
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .helpers import add_centered, add_body, make_table


def build(doc):
    doc.add_heading(
        "CHƯƠNG 5. CÀI ĐẶT, TRIỂN KHAI VÀ KẾT QUẢ THỰC HIỆN", level=1
    )

    # ── 5.1 ──────────────────────────────────
    doc.add_heading("5.1. Mô hình triển khai ứng dụng", level=2)

    add_body(
        doc,
        "Hệ thống SmartFridge AI được triển khai bằng Docker Compose gồm 4 "
        "container dịch vụ chạy trên cùng một mạng nội bộ (fridge-network). "
        "Mô hình này cho phép các dịch vụ giao tiếp với nhau thông qua tên "
        "DNS nội bộ mà không cần cấu hình IP tĩnh."
    )
    add_body(doc, "Các dịch vụ trong docker-compose.yml:")
    add_body(
        doc,
        "(1) postgres (pgvector:pg15): Cơ sở dữ liệu PostgreSQL với phần mở "
        "rộng pgvector, chạy trên port 5432. Có health check bằng pg_isready."
    )
    add_body(
        doc,
        "(2) backend (Spring Boot): Ứng dụng backend Java, chạy trên port "
        "8080. Health check thông qua /actuator/health. Phụ thuộc vào "
        "postgres khởi động thành công."
    )
    add_body(
        doc,
        "(3) redis (Redis 7): Bộ nhớ đệm, chạy trên port 6379. Health check "
        "bằng redis-cli ping."
    )
    add_body(
        doc,
        "(4) ai_service (FastAPI): Dịch vụ AI Python, chạy trên port 8001. "
        "Health check thông qua /health. Phụ thuộc vào postgres, redis và "
        "backend đều sẵn sàng."
    )
    add_body(
        doc,
        "Ứng dụng mobile (React Native) được build và chạy trực tiếp trên "
        "thiết bị thật hoặc emulator, kết nối đến backend thông qua địa chỉ "
        "IP của máy chủ (hoặc localhost khi chạy cùng máy)."
    )

    add_centered(doc,
                 "[Hình 5.1. Mô hình triển khai Docker Compose]",
                 size=12, bold=True)

    # ── 5.2 ──────────────────────────────────
    doc.add_heading("5.2. Các bước cài đặt và triển khai", level=2)

    add_body(doc, "Bước 1 — Clone mã nguồn:", bold=True)
    add_body(doc, "git clone https://github.com/phucmouse135/mobile.git")
    add_body(doc, "cd mobile")

    add_body(doc, "Bước 2 — Cấu hình môi trường:", bold=True)
    add_body(
        doc,
        "Tạo file .env cho backend với các thông số: DATABASE_URL, "
        "JWT_SECRET, GOOGLE_CLIENT_ID, GOOGLE_CLIENT_SECRET. Tạo file .env "
        "cho AI service với: GEMINI_API_KEY, DATABASE_URL, REDIS_URL, "
        "OPENWEATHER_API_KEY. Cấu hình EXPO_PUBLIC_API_URL trong file .env "
        "của mobile-app trỏ đến địa chỉ backend."
    )

    add_body(doc, "Bước 3 — Khởi động hệ thống bằng Docker:", bold=True)
    add_body(doc, "docker compose up --build -d")
    add_body(
        doc,
        "Lệnh này sẽ tự động: pull image PostgreSQL và Redis, build backend "
        "và AI service từ Dockerfile, chạy Flyway migration tạo bảng, và "
        "khởi động tất cả dịch vụ theo đúng thứ tự phụ thuộc."
    )

    add_body(doc, "Bước 4 — Cài đặt và chạy mobile:", bold=True)
    add_body(doc, "cd mobile-app")
    add_body(doc, "npm install")
    add_body(doc, "npx react-native run-android  (hoặc run-ios)")

    add_body(doc, "Bước 5 — Kiểm tra hoạt động:", bold=True)
    add_body(
        doc,
        "Truy cập http://localhost:8080/actuator/health để kiểm tra backend. "
        "Truy cập http://localhost:8001/health để kiểm tra AI service. Mở "
        "ứng dụng trên điện thoại, đăng ký tài khoản và kiểm tra các chức năng."
    )

    # ── 5.3 ──────────────────────────────────
    doc.add_heading("5.3. Kết quả thực hiện được", level=2)

    add_body(
        doc,
        "Phần này trình bày kết quả thực hiện các chức năng chính của hệ "
        "thống, tập trung vào phần của Nguyễn Đình Phúc."
    )

    add_body(doc, "Chức năng đăng nhập và xác thực:", bold=True)
    add_body(
        doc,
        "Hệ thống đăng nhập hoạt động ổn định với cả hai phương thức: "
        "email/mật khẩu và Google OAuth2. JWT token được tạo thành công "
        "với access token (30 phút) và refresh token (30 ngày). Ứng dụng "
        "tự động refresh token khi hết hạn mà không cần người dùng đăng "
        "nhập lại. Màn hình Login có validation đầy đủ và hiển thị thông "
        "báo lỗi rõ ràng bằng tiếng Việt."
    )

    add_centered(doc,
                 "[Hình 5.2. Kết quả giao diện màn hình Login]",
                 size=12, bold=True)

    add_body(doc, "Chức năng quản lý Household:", bold=True)
    add_body(
        doc,
        "Tạo Household thành công với đầy đủ thông tin: tên nhà, mã mời "
        "6 ký tự và QR code. Thành viên có thể tham gia nhanh bằng cách "
        "quét QR hoặc nhập mã thủ công. Phân quyền Owner/Admin/Member hoạt "
        "động chính xác. Owner có thể chuyển quyền sở hữu, Admin có thể "
        "mời/xóa thành viên. Chức năng tái tạo mã mời (regenerate) vô hiệu "
        "mã cũ và tạo mã mới thành công."
    )

    add_body(doc, "Chức năng quản lý kho thực phẩm (Inventory):", bold=True)
    add_body(
        doc,
        "Danh sách thực phẩm hiển thị đầy đủ thông tin: tên, số lượng, đơn "
        "vị, vị trí lưu trữ và trạng thái hạn sử dụng. Thao tác CRUD hoạt "
        "động ổn định. Thuật toán FIFO trừ kho chính xác — ưu tiên lô cũ "
        "nhất trước. Hệ thống tự động cảnh báo khi thực phẩm còn dưới "
        "3 ngày đến hạn."
    )

    add_centered(doc,
                 "[Hình 5.3. Kết quả giao diện màn hình Inventory]",
                 size=12, bold=True)

    add_body(doc, "Chức năng đồng bộ thời gian thực:", bold=True)
    add_body(
        doc,
        "WebSocket/STOMP hoạt động ổn định giữa nhiều thiết bị. Khi một "
        "thiết bị cập nhật kho, các thiết bị khác nhận được thay đổi trong "
        "vòng dưới 1 giây. Cơ chế reconnect tự động khi mất kết nối đảm "
        "bảo dữ liệu luôn đồng bộ."
    )

    add_centered(doc,
                 "[Hình 5.4. Kết quả giao diện Household và QR]",
                 size=12, bold=True)

    # ── 5.4 ──────────────────────────────────
    doc.add_heading("5.4. Kết quả thử nghiệm/triển khai", level=2)

    add_body(
        doc,
        "Nhóm đã thực hiện kiểm thử hệ thống với các kết quả được tổng "
        "hợp trong bảng sau:"
    )

    make_table(doc, ["Chỉ tiêu", "Kết quả", "Ghi chú"], [
        ["Tổng số module hoàn thành", "8/8",
         "Auth, Household, Inventory, Recipe,\nPlanner, Chat, Social, Notification"],
        ["Số màn hình giao diện", "20+",
         "Bao gồm tất cả màn hình trong\n5 tab chính và các màn hình con"],
        ["Số API endpoint tích hợp", "30+",
         "REST API (backend) + 1 AI suggestion endpoint"],
        ["Số bảng cơ sở dữ liệu", "15+",
         "Được tạo tự động bởi 21 Flyway migrations"],
        ["Số loại thực phẩm test trong kho", "35",
         "Bao gồm rau củ, thịt cá, sữa,\ntrái cây, gia vị, đồ đông lạnh"],
        ["Số user/household test mẫu", "8 user / 3 household",
         "Mỗi household có 2–4 thành viên"],
        ["Số luồng nghiệp vụ test thành công", "25/25",
         "Bao gồm tất cả CRUD, FIFO, sync,\nauth, invite flow"],
        ["Số tình huống sync kiểm thử", "10",
         "Bao gồm: thêm/sửa/xóa, đa thiết bị,\nmất kết nối, reconnect"],
        ["Thời gian phản hồi trung bình API", "< 500ms",
         "Đo trên môi trường Docker local"],
        ["Thời gian đồng bộ WebSocket", "< 1 giây",
         "Đo giữa 2 thiết bị Android\ntrong cùng mạng LAN"],
    ])

    # ── 5.5 ──────────────────────────────────
    doc.add_heading("5.5. Đánh giá kết quả", level=2)

    add_body(
        doc,
        "Nhìn chung, hệ thống SmartFridge AI đã đạt được các mục tiêu đề "
        "ra ban đầu. Các điểm nổi bật bao gồm:"
    )
    add_body(
        doc,
        "Về mặt chức năng: Hệ thống hoàn thành đầy đủ 8 nhóm module chính "
        "từ xác thực, quản lý kho, gợi ý món ăn, kế hoạch bữa ăn cho đến "
        "chatbot và thông báo. Các luồng nghiệp vụ hoạt động ổn định và "
        "nhất quán giữa các thành phần."
    )
    add_body(
        doc,
        "Về mặt kỹ thuật: Kiến trúc microservice với Docker Compose cho "
        "phép triển khai nhanh và dễ bảo trì. Việc tách biệt AI service "
        "khỏi backend chính là quyết định thiết kế đúng đắn, giúp hai "
        "thành phần có thể scale độc lập. WebSocket/STOMP đảm bảo đồng "
        "bộ thời gian thực hiệu quả."
    )
    add_body(
        doc,
        "Về mặt phân công: Mô hình Vertical Slice đã chứng minh hiệu quả "
        "trong việc phân công công việc cho nhóm 4 thành viên. Mỗi người "
        "chịu trách nhiệm toàn bộ \"lát dọc\" của mình, từ giao diện đến "
        "cơ sở dữ liệu, giúp giảm xung đột code và tăng tính chủ động."
    )
    add_body(
        doc,
        "Phần của Nguyễn Đình Phúc (Core & Sync) đóng vai trò nền tảng cho "
        "toàn bộ hệ thống. Module xác thực và Household là điều kiện tiên "
        "quyết để các module khác hoạt động. Đồng bộ thời gian thực bằng "
        "WebSocket là tính năng kỹ thuật có độ phức tạp cao nhất trong hệ "
        "thống và đã được triển khai thành công."
    )

    doc.add_page_break()
